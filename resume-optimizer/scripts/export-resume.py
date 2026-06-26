#!/usr/bin/env python3
"""将 Markdown 简历导出为 DOCX 格式。"""

import sys
import os

def export_to_docx(md_path: str, docx_path: str):
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("错误: 需要安装 python-docx，请运行: pip install python-docx", file=sys.stderr)
        sys.exit(1)

    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    doc = Document()

    # 设置默认字体
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Microsoft YaHei"
    font.size = Pt(10.5)

    for line in lines:
        line = line.rstrip("\n")

        if not line.strip():
            continue

        # 标题
        if line.startswith("# "):
            p = doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            p = doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            p = doc.add_heading(line[4:], level=3)
        elif line.startswith("---"):
            continue
        elif line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(line[2:], style="List Bullet")
        else:
            p = doc.add_paragraph(line)

    doc.save(docx_path)
    print(f"✅ 已导出: {docx_path}")

def main():
    if len(sys.argv) < 3:
        print("用法: python export-resume.py <输入markdown文件> <输出docx文件>", file=sys.stderr)
        sys.exit(1)

    md_path = sys.argv[1]
    docx_path = sys.argv[2]

    if not os.path.exists(md_path):
        print(f"错误: 文件不存在: {md_path}", file=sys.stderr)
        sys.exit(1)

    export_to_docx(md_path, docx_path)

if __name__ == "__main__":
    main()
